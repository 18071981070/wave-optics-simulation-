from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_DIR = Path("竞赛材料")
OUTPUT_PATH = OUTPUT_DIR / "波动光学仿真平台_软件结构代码来源与AI辅助使用说明.docx"

FONT_NAME = "SimSun"
BODY_SIZE = 12
BLACK = "000000"
NAVY = BLACK
BLUE = BLACK
DARK = BLACK
MUTED = BLACK
LINE = "D8E1EA"
PALE_BLUE = "EAF1F8"
PALE_TEAL = "EAF7F5"
PALE_GOLD = "FFF6DF"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name=FONT_NAME, size=BODY_SIZE, bold=None, color=BLACK, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_runs(paragraph, parts):
    for text, opts in parts:
        run = paragraph.add_run(text)
        set_font(run, **opts)
    return paragraph


def add_body(doc, text="", bold_prefix=None, after=5, first_indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(21)
    if bold_prefix and text.startswith(bold_prefix):
        add_runs(p, [
            (bold_prefix, {"bold": True, "color": NAVY}),
            (text[len(bold_prefix):], {}),
        ])
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text, level=0, after=3):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    set_font(p.add_run(text), size=BODY_SIZE)
    return p


def add_number(doc, text, level=0, after=3):
    p = doc.add_paragraph(style="List Number" if level == 0 else "List Number 2")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    set_font(p.add_run(text), size=BODY_SIZE)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    return p


def add_note(doc, label, text, fill=PALE_GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    add_runs(p, [
        (label + " ", {"bold": True, "color": NAVY}),
        (text, {"size": 10.2}),
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=PALE_BLUE, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), size=BODY_SIZE, bold=True, color=BLACK)
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    header_tr_pr.append(tbl_header)
    for row in rows:
        cells = table.add_row().cells
        tr_pr = table.rows[-1]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_font(p.add_run(str(value)), size=BODY_SIZE, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=BODY_SIZE, color=BLACK)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    tail = paragraph.add_run("页")
    set_font(tail, size=BODY_SIZE, color=BLACK)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLACK, 14, 7),
        (2, 14, BLACK, 10, 5),
        (3, 12, BLACK, 7, 3),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Bullet 2", "List Number", "List Number 2"]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.size = Pt(BODY_SIZE)
        style.font.color.rgb = RGBColor.from_string(BLACK)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("软件结构、代码来源与 AI 辅助使用说明"), size=BODY_SIZE, color=BLACK)
    add_page_number(section.footer.paragraphs[0])


def normalize_all_text(doc):
    """Final typography pass: SimSun, black text, body/table size = 小四."""
    def normalize_paragraph(paragraph, force_body=True):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_font(
                run,
                name=FONT_NAME,
                size=None if is_heading and not force_body else BODY_SIZE,
                color=BLACK,
                bold=run.bold,
                italic=run.italic,
            )
        if not is_heading:
            paragraph.paragraph_format.line_spacing = 1.25

    for paragraph in doc.paragraphs:
        normalize_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    normalize_paragraph(paragraph)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            normalize_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            normalize_paragraph(paragraph)


def build_document():
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)

    # Memo-style title block; no identifying institution or person information.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("技术说明文档"), size=BODY_SIZE, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run("波动光学交互式仿真平台"), size=18, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run("软件整体结构、代码来源与 AI 辅助使用说明"), size=BODY_SIZE, color=BLACK)

    meta = add_table(
        doc,
        ["文档用途", "适用范围", "版本日期"],
        [["竞赛参赛文档附件", "当前项目代码与资源文件", "2026年8月22日"]],
        [2500, 4360, 2500],
        header_fill=PALE_TEAL,
        font_size=9.5,
    )
    for cell in meta.rows[1].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_note(
        doc,
        "提交说明：",
        "本文档依据当前代码仓库整理，用于响应“说明软件整体结构，明确调用的已有函数、引用代码与自写代码”的要求。文档未写入学校、学院、指导教师或学生姓名等信息。",
        fill=PALE_GOLD,
    )

    add_heading(doc, "1. 说明目的与标记规则", 1)
    add_body(doc, "本文档对作品的软件架构、主要模块、第三方依赖、项目自写函数、外部资料引用以及 AI 辅助情况进行可追溯说明。为便于评审查验，全文采用以下标记：")
    add_table(
        doc,
        ["标记", "含义", "使用场景"],
        [
            ("[自写]", "针对本作品业务逻辑编写的项目代码", "光学模型、参数联动、页面流程、本地问答等"),
            ("[调用]", "调用第三方库或 Python 标准库的公开接口", "Streamlit、NumPy、Matplotlib、Pandas、Pillow、Requests 等"),
            ("[引用]", "引用的理论公式、公开文档、外部资源或第三方代码", "文献公式、官方 API 文档、外部图片等"),
            ("[AI辅助]", "生成式 AI 参与建议、调试、改写或文档整理，并经人工复核", "界面样式、公式检查、异常修复、提示词与本文档"),
        ],
        [1500, 3650, 4210],
    )

    add_heading(doc, "2. 软件整体结构", 1)
    add_body(doc, "作品采用“交互界面层—物理计算层—智能辅助层—教学内容层—资源与配置层”的分层结构。其中，实验模型、参数体系、交互流程、图表组合和教学任务均由项目组自主设计并编写。用户在 Streamlit 页面调节参数，主程序将参数传入项目自写的光学计算模块，再用 NumPy 数组和 Matplotlib 公开绘图接口返回仿真图样、强度曲线及定量指标；智能助手根据当前实验和参数组装上下文，通过系统环境变量中的 API 配置请求大模型，失败时使用项目自写的本地物理知识回答。")

    add_heading(doc, "2.1 项目组自主设计与编写范围", 2)
    add_note(doc, "自主开发确认：", "本作品的核心业务不是将某个第三方软件原样拼接为页面，而是由项目组自主完成实验定义、物理模型、数据结构、交互联动和教学流程的设计。下列内容属于 [自写]，不属于第三方代码复制：")
    for item in [
        "六类实验的参数范围、单位、默认值和实验任务设计。",
        "双缝、单缝、光栅、迈克耳孙、薄膜和偏振模型的数值实现，以及光强、相位、主极大、中央明纹宽度、可见度等特征量的联动。",
        "仪器状态—观察屏—强度曲线—数据证据的全宽交互展示逻辑。",
        "学生学习端、教师教学端、教学模式、练习模式、探究记录和数据导出流程。",
        "智能助手的实验上下文、快捷问题、本地降级回答、异常保护和公式 LaTeX 渲染。",
        "可视化色图、装置图统一画布、大屏观察屏和页面风格设计。",
    ]:
        add_bullet(doc, item, after=3)
    add_table(
        doc,
        ["层级", "主要文件", "职责", "来源标记"],
        [
            ("交互界面层", "wave_optics_simulation.py", "页面布局、学生/教师端、参数控件、图表、数据导出、智能助手入口", "[自写][调用][AI辅助]"),
            ("物理计算层", "optical_calculator.py", "六类波动光学实验、误差模型、反推计算与特征量", "[自写][引用][AI辅助]"),
            ("智能辅助层", "agent_module_v2.py", "API 配置、上下文组装、对话、本地降级、快捷问题", "[自写][调用][AI辅助]"),
            ("教学内容层", "teaching_module.py、exercise_module.py", "实验原理、步骤、提示、分层练习题", "[自写][引用]"),
            ("资源与配置层", "requirements.txt、simhei.ttf、图片资源、run_simulation.bat", "依赖、中文字体、光路图与启动入口", "[调用][引用]/项目配置"),
        ],
        [1420, 2350, 3590, 2000],
        font_size=8.8,
    )

    add_heading(doc, "3. 第三方库与已有函数调用说明", 1)
    add_body(doc, "下列内容属于对公开软件库接口的正常调用，不等同于复制该软件库的源代码。版本范围以 requirements.txt 和实际运行环境为准。")
    add_table(
        doc,
        ["依赖/来源", "主要调用的已有接口", "在本作品中的用途", "标记"],
        [
            ("Streamlit [3]", "st.set_page_config、st.slider、st.radio、st.selectbox、st.pyplot、st.metric、st.chat_message、st.session_state", "网页交互、会话状态、图像展示与聊天组件", "[调用]"),
            ("NumPy [4]", "np.linspace、np.sin、np.cos、np.sinc、np.meshgrid、np.clip、np.unwrap、np.random", "采样、数值运算、干涉/衍射数组、噪声模拟", "[调用]"),
            ("Matplotlib [5]", "plt.figure、plt.subplots、Axes.plot、Axes.imshow、fill_between、LinearSegmentedColormap", "绘制观察屏、光强曲线、偏振椭圆和同心环", "[调用]"),
            ("Pandas/OpenPyXL [6]", "pd.DataFrame、pd.ExcelWriter、DataFrame.to_excel", "仿真数据表格化和 Excel 导出", "[调用]"),
            ("Pillow [7]", "Image.open、convert、thumbnail、Image.new、paste", "实验装置图统一尺寸、缩放和画布合成", "[调用]"),
            ("Requests [8]", "requests.post、requests.get、Timeout、ConnectionError", "调用 DashScope/OpenAI 兼容聊天接口或本地 Ollama 接口", "[调用]"),
            ("Python 标准库", "os.getenv、BytesIO、re.search、typing", "环境变量、内存文件、数值提取和类型标注", "[调用]"),
            ("DashScope 兼容接口 [9]", "/compatible-mode/v1/chat/completions", "使用系统环境变量中的 API Key 进行智能问答", "[调用]"),
        ],
        [1600, 3100, 3460, 1200],
        font_size=8.45,
    )

    add_heading(doc, "4. 项目自写代码与函数清单", 1)
    add_body(doc, "“自写”表示该函数为本项目组根据实验需求自主定义并编写的业务函数，而不是第三方库中的现成函数。其内部可以合法调用 NumPy、Matplotlib 等公开 API，但计算流程、参数检查、特征量输出和页面联动均由项目组自主编写。本轮界面、公式与智能助手修改包含 AI 辅助，但不改变项目组对核心代码和最终成果的责任。见第 6 节。")

    add_heading(doc, "4.4 项目组自主开发工作量概括", 2)
    add_table(
        doc,
        ["自主开发项目", "具体工作", "提交标记"],
        [
            ("实验模型自主建模", "将六类波动光学实验转化为可调参数、可观察图样、可验证特征量的计算模型，并为每个模型设计默认参数。", "[自写]"),
            ("参数—图像—数据联动", "实现滑块变化后仪器展示、观察屏、光强曲线和指标同步变化，形成实验特性的可视化闭环。", "[自写]"),
            ("教学端功能自主设计", "将原理、公式、参数操作、证据解释、教师脚本和练习任务组织为一条学习路径。", "[自写]"),
            ("非理想因素与评估", "加入环境光、随机噪声、误差统计、测量记录和数据导出，将理想模型扩展为可教学评估的虚拟实验。", "[自写]"),
            ("智能助手工程化处理", "自主设计环境变量配置、上下文传递、降级回答、快捷提问、异常保护和公式渲染。", "[自写][调用]"),
        ],
        [2200, 5100, 2260],
        font_size=8.65,
    )

    add_heading(doc, "4.1 物理计算模块（optical_calculator.py）", 2)
    add_table(
        doc,
        ["函数/方法", "功能说明", "来源标记"],
        [
            ("double_slit_interference", "双缝干涉、有限缝宽包络与部分相干性计算", "[自写][引用][AI辅助]"),
            ("single_slit_diffraction", "单缝夫琅禾费衍射光强及暗纹/中央明纹特征", "[自写][引用][AI辅助]"),
            ("multi_slit_diffraction", "多缝干涉因子与单缝包络耦合、主极大与分辨本领", "[自写][引用][AI辅助]"),
            ("michelson_interferometer", "迈克耳孙光程差、光强比、可见度和条纹移动数", "[自写][引用][AI辅助]"),
            ("thin_film_interference", "等厚/等倾干涉、折射角、反射相位突变与强度", "[自写][引用][AI辅助]"),
            ("polarization_interference", "旋转矩阵、波片相位延迟、分析器投影与偏振态", "[自写][引用][AI辅助]"),
            ("add_noise / calculate_error", "环境光、探测器/随机噪声及误差统计", "[自写]"),
            ("verify_double_slit / young_double_slit", "双缝模型验证与兼容计算入口", "[自写]"),
            ("airy_disk / fresnel_diffraction", "圆孔艾里斑与菲涅耳衍射扩展模型", "[自写][引用]"),
            ("calculate_wavelength_from_fringe", "由条纹间距、缝距和屏距反推波长并进行不确定度传播", "[自写][引用]"),
            ("generate_wavelength_color", "将可见光波长区间转为中文颜色名称", "[自写]"),
        ],
        [2650, 4450, 2260],
        font_size=8.65,
    )

    add_heading(doc, "4.2 主界面与可视化（wave_optics_simulation.py）", 2)
    add_table(
        doc,
        ["函数/代码区", "功能说明", "来源标记"],
        [
            ("wavelength_to_css / wavelength_to_rgb", "将波长映射为界面光束颜色和 Matplotlib RGB 颜色", "[自写][AI辅助]"),
            ("optical_cmap", "构建从黑色到单色光的自定义色图", "[自写][调用]"),
            ("load_uniform_diagram", "将不同尺寸的光路示意图归一化到统一画布", "[自写][调用][AI辅助]"),
            ("六类实验参数组与图表组装", "连接滑块、计算函数、仪器状态、观察屏、强度曲线和特征量", "[自写][调用][AI辅助]"),
            ("submit_agent_question / clear_agent_conversation", "智能助手发送、快捷问答、清空与 Streamlit 回调", "[自写][AI辅助]"),
            ("CSS 主题与响应式布局", "实验控制台风格、文字对比度、无白闪重算、大屏/移动端布局", "[自写][AI辅助]"),
        ],
        [2750, 4300, 2310],
        font_size=8.7,
    )

    add_heading(doc, "4.3 智能助手（agent_module_v2.py）", 2)
    add_table(
        doc,
        ["函数组", "功能说明", "来源标记"],
        [
            ("环境配置：_load_environment_config、refresh_environment_config", "读取 DASHSCOPE_API_KEY 等系统环境变量，不在界面显示密钥", "[自写][AI辅助]"),
            ("上下文：set_experiment_context、get_system_prompt", "将实验名、当前参数和心要公式注入提示词", "[自写][AI辅助]"),
            ("API：test_api_connection、call_api、_call_dashscope_api、_call_ollama_api", "构造 HTTP 请求、解析回复、处理超时和连接异常", "[自写][调用][AI辅助]"),
            ("对话：generate_response、add_message", "管理消息历史、传递当前问题、防止 API 异常导致页面崩溃", "[自写][AI辅助]"),
            ("降级回答：get_local_response、_get_physics_response、_get_general_response", "远程模型不可用时，根据实验参数给出公式、计算或原理回答", "[自写][引用][AI辅助]"),
            ("get_quick_questions / suggest_parameters", "按当前实验生成快捷问题和参数建议", "[自写][AI辅助]"),
        ],
        [2900, 4210, 2250],
        font_size=8.65,
    )

    add_heading(doc, "5. 引用代码、理论公式与外部资源说明", 1)
    add_heading(doc, "5.1 引用代码判定", 2)
    add_body(doc, "经项目组对本次提交版代码文件的核查，未发现将某个第三方项目的整段源代码、完整函数文件或开源项目源码直接复制进本项目的情况。本提交版的核心计算函数、参数体系、页面交互、教学模式、练习内容和智能助手业务逻辑由项目组自主开发。对 Streamlit、NumPy、Matplotlib 等属于对公开软件包接口的 [调用]，不属于“引用代码”。若历史开发过程中存在从教程、博客或开源项目改写的代码片段，项目组仍应在最终提交前补充原始链接、作者、许可证和改写位置。")
    add_note(doc, "需项目组复核：", "图片资源文件的原始作者/生成方式与授权情况无法仅由文件名判定。若非项目组自制，必须在图注或资源清单中补充“来源链接、原作者、授权方式”；若为 AI 生成或 AI 修图，则标记 [AI辅助]。")

    add_heading(doc, "5.2 理论公式与知识来源", 2)
    add_body(doc, "双缝干涉、夫琅禾费衍射、光栅方程、迈克耳孙干涉、薄膜干涉、马吕斯定律、琼斯矩阵等为经典波动光学理论，应在研究报告的“理论基础”和 PPT 对应页面标注参考文献 [1]-[2]。数值实现和软件接口的使用可标注官方文桮 [3]-[9]。")
    add_table(
        doc,
        ["内容", "建议页面标注", "标记"],
        [
            ("干涉、衍射、偏振、薄膜公式", "“理论公式参考：[1]-[2]”", "[引用]"),
            ("Streamlit 界面组件", "“软件框架参考 Streamlit 官方文档 [3]”", "[调用][引用]"),
            ("数值计算和绘图", "“数值/绘图 API 参考 [4]-[5]”", "[调用][引用]"),
            ("AI 连接与对话", "“模型接口参考 [8]-[9]；提示词与业务逻辑为项目实现”", "[调用][AI辅助]"),
        ],
        [3000, 4410, 1950],
        font_size=8.9,
    )

    add_heading(doc, "6. AI 辅助使用情况说明", 1)
    add_body(doc, "本作品开发与文档整理过程中使用了生成式 AI 作为辅助工具。AI 不代替项目组对物理模型、公式、代码和参赛材料的最终审核责任；所有修改均应经过人工阅读、运行检查和结果比对。")
    add_table(
        doc,
        ["AI 辅助内容", "使用方式", "人工复核/验证", "标记位置"],
        [
            ("界面与交互优化", "提供版式、CSS、模块顺序、图像比例与文字可读性改进建议", "人工选择并运行 Streamlit 页面检查", "wave_optics_simulation.py 相关样式与布局 [AI辅助]"),
            ("物理模型与公式检查", "辅助检查量纲、极值条件、特征量和参数单位的实现", "与经典光学公式、默认参数和数值结果比对", "optical_calculator.py 修订部分 [AI辅助]"),
            ("智能助手调试", "辅助修复提示词转义、当前问题遗漏、快捷提问回调、异常降级和 LaTeX 渲染", "语法编译、六类实验运行测试、快捷/手动问答测试", "agent_module_v2.py 及聊天界面 [AI辅助]"),
            ("本说明文档", "辅助盘点代码文件、组织结构、生成 Word 排版初稿", "基于实际仓库逐项核对函数名、依赖和模块职责", "文档扉页及第 6 节 [AI辅助]"),
        ],
        [1900, 3250, 2860, 1350],
        font_size=8.45,
    )

    add_heading(doc, "6.1 可直接放入研究报告/PPT 的 AI 声明", 2)
    add_note(
        doc,
        "[AI辅助声明]",
        "本作品在界面样式优化、部分代码调试、物理公式实现检查、智能助手提示词与异常处理、以及技术说明文档整理中使用了生成式 AI 辅助。AI 输出均经项目组阅读、修改、运行测试和结果复核；实验方案、物理模型选择、核心功能取舍和最终参赛成果由项目组负责。",
        fill=PALE_TEAL,
    )

    add_heading(doc, "7. 代码片段的规范标注示例", 1)
    add_body(doc, "建议在研究报告或附录中对关键代码采用如下写法：")
    examples = [
        ("[调用] Streamlit 控件", 'wavelength = st.slider("波长 (nm)", 400, 760, 540, 10)'),
        ("[自写] 项目物理函数", 'x, intensity, phase, info = calculator.double_slit_interference(...)'),
        ("[调用] 第三方 HTTP 接口", 'response = requests.post(api_url, headers=headers, json=payload, timeout=60)'),
        ("[AI辅助] 对话稳定性修改", 'content = content.replace(r"\\[", "$$").replace(r"\\]", "$$")'),
    ]
    for label, code in examples:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(label), size=9.3, bold=True, color=NAVY)
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Pt(14)
        cp.paragraph_format.right_indent = Pt(8)
        cp.paragraph_format.space_after = Pt(6)
        run = cp.add_run(code)
        set_font(run, name="Consolas", size=8.5, color=DARK)
        p_pr = cp._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F4F6F8")
        p_pr.append(shd)

    add_heading(doc, "8. 提交前合规核对清单", 1)
    checks = [
        "研究报告中的光学公式和理论结论已标注参考文献 [1]-[2]。",
        "PPT 中出现公式、结论或外部图像的页面已就地标注来源。",
        "所有非自制图片、音频、视频、PPT 素材已记录原链接、原作者和授权状态。",
        "代码中如有从教程/开源项目改写的片段，已标注文件、行段、原链接和许可证。",
        "研究报告、PPT 和视频中已加入第 6.1 节的 AI 辅助声明。",
        "文档和展示材料中未出现学校、学院、指导教师或学生姓名等可影响公平评审的信息。",
        "六类实验、数据导出、快捷提问、手动提问和清空对话均已在最终版本上运行检查。",
    ]
    for item in checks:
        add_bullet(doc, "☐ " + item, after=4)

    add_heading(doc, "9. 参考文献与官方资料", 1)
    refs = [
        "[1] Eugene Hecht. Optics, 5th Edition. Pearson, 2017.",
        "[2] Max Born, Emil Wolf. Principles of Optics, 7th (Expanded) Edition. Cambridge University Press, 1999.",
        "[3] Streamlit Documentation. https://docs.streamlit.io/ （访问日期：2026-08-22）。",
        "[4] NumPy Documentation. https://numpy.org/doc/stable/ （访问日期：2026-08-22）。",
        "[5] Matplotlib Documentation. https://matplotlib.org/stable/ （访问日期：2026-08-22）。",
        "[6] pandas Documentation. https://pandas.pydata.org/docs/ （访问日期：2026-08-22）。",
        "[7] Pillow Documentation. https://pillow.readthedocs.io/ （访问日期：2026-08-22）。",
        "[8] Requests Documentation. https://requests.readthedocs.io/ （访问日期：2026-08-22）。",
        "[9] 阿里云模型服务灵积 DashScope 官方文档. https://help.aliyun.com/zh/model-studio/ （访问日期：2026-08-22）。",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(ref), size=9.3)

    add_note(
        doc,
        "结论：",
        "本项目的主要业务函数和界面流程属项目代码 [自写]；第三方软件通过公开 API [调用]；经典光学理论和官方文档已在参考文献中 [引用]；界面优化、部分公式实现检查、智能助手调试和本文档整理已如实标注 [AI辅助]。",
        fill=PALE_TEAL,
    )

    normalize_all_text(doc)

    props = doc.core_properties
    props.title = "波动光学仿真平台软件结构、代码来源与AI辅助使用说明"
    props.subject = "竞赛参赛文档附件"
    props.author = "参赛项目组"
    props.keywords = "软件结构, 代码来源, AI辅助, 波动光学"
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH.resolve())


if __name__ == "__main__":
    build_document()
